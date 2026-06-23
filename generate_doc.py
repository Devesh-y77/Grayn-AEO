import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import urllib.request
import base64

# --- 1. Generate Diagram ---
mermaid_code = """
graph TD;
    Client[Frontend Dashboard] -->|HTTPS| FastAPI[FastAPI Backend - Railway];
    SlackApp[Slack Custom App] -->|Webhooks / Slash Commands| FastAPI;
    
    FastAPI -->|Reads & Writes| Supabase[(Supabase PostgreSQL)];
    
    FastAPI -->|Schedules via APScheduler| TrackingService[Live Tracking Engine];
    TrackingService -->|Natively Scrapes| Playwright[Playwright Browser Automation];
    TrackingService -->|Dispatches Prompts| AI_Router[AI Provider Service];
    
    AI_Router --> OpenAI[OpenAI API];
    AI_Router --> Gemini[Google Gemini API];
    AI_Router --> Perplexity[Perplexity API];
    AI_Router --> Claude[Anthropic API];
    
    FastAPI --> ContentAnalyzer[Content Gaps Service];
    ContentAnalyzer -->|Scrapes Top 15 URLs| Playwright;
    ContentAnalyzer -->|Identifies Semantic Gaps| AI_Router;
"""
try:
    encoded = base64.b64encode(mermaid_code.encode('utf-8')).decode('utf-8')
    url = f"https://mermaid.ink/img/{encoded}?theme=neutral"
    req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
    with urllib.request.urlopen(req) as response, open("backend_diagram.png", 'wb') as out_file:
        out_file.write(response.read())
    has_diagram = True
except Exception as e:
    print("Failed to download diagram:", e)
    has_diagram = False

# --- 2. Create Document ---
doc = docx.Document()

# Styles
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('Grayn AEO: Comprehensive Technical Architecture & Integration Guide', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("This document serves as the absolute source of truth for the Grayn Answer Engine Optimization (AEO) backend. It covers the architectural flow, database connection protocols, Railway deployment lifecycle, and an exhaustive guide on integrating the system with Slack.")

# --- DIAGRAM ---
doc.add_heading('1. Backend Network Architecture Flowchart', level=1)
doc.add_paragraph("The following diagram illustrates how the client frontend, Slack applications, database, and third-party AI APIs connect through our central Railway-hosted FastAPI infrastructure.")

if has_diagram:
    doc.add_picture("backend_diagram.png", width=Inches(6.0))
    last_paragraph = doc.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
else:
    doc.add_paragraph("[Diagram generation failed. Please see the codebase for architecture details.]")

# --- BACKEND STRUCTURE ---
doc.add_heading('2. In-Depth Backend Structure', level=1)

doc.add_heading('A. The FastAPI Core (app/main.py)', level=2)
doc.add_paragraph("The heart of the application is a highly asynchronous FastAPI server. It is responsible for handling all incoming HTTP requests. Upon startup, it initializes the CORS middleware, connects to the Supabase client, and spins up the APScheduler which manages background tracking batches.")

doc.add_heading('B. Routing Layer (app/routers/)', level=2)
p = doc.add_paragraph()
p.add_run("• /internal.py: ").bold = True
p.add_run("These routes are consumed specifically by the Next.js frontend dashboard (e.g., creating workspaces, managing billing, fetching detailed charts).\n")
p.add_run("• /v1.py: ").bold = True
p.add_run("This is the public-facing API version 1. It contains endpoints for external integrations, including the dedicated `/v1/slack/` webhook routes.")

doc.add_heading('C. Business Logic Services (app/services/)', level=2)
p = doc.add_paragraph()
p.add_run("• tracking.py: ").bold = True
p.add_run("When a user triggers a tracking run, this service dispatches the target keyword to up to 7 different AI engines concurrently. It uses Playwright to navigate sites where API access isn't sufficient.\n")
p.add_run("• content_analyzer.py: ").bold = True
p.add_run("Powers the 'Content Gaps Studio'. It scrapes the top 15 domains that the AIs cited, aggregates their text, and feeds it into OpenAI (gpt-4o) or Gemini to generate Markdown strategic briefs.\n")
p.add_run("• scoring.py: ").bold = True
p.add_run("Houses the mathematical models for calculating 'AI Visibility Percentage' and 'Share of Voice' across competitors.")

# --- DEPLOYMENT & CONNECTIONS ---
doc.add_heading('3. Deployment Lifecycle & Connection Strings', level=1)

doc.add_heading('A. Deploying to Railway', level=2)
doc.add_paragraph("The backend is containerized using Docker and is built to be deployed on Railway.app. To deploy:")
doc.add_paragraph("1. Connect the GitHub repository to a new Railway project.", style='List Number')
doc.add_paragraph("2. Railway will automatically detect the Dockerfile in the /backend directory.", style='List Number')
doc.add_paragraph("3. In the Railway project settings, map the internal PORT (usually 8000) and generate a Public Domain (e.g., grayn-api.railway.app).", style='List Number')

doc.add_heading('B. Critical Environment Variables', level=2)
doc.add_paragraph("You must define the following variables inside the Railway 'Variables' tab. The application will crash or severely limit functionality if these are missing.")

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Variable Name'
hdr_cells[1].text = 'Purpose'
hdr_cells[2].text = 'Example / Format'

def add_row(var, purp, ex):
    row_cells = table.add_row().cells
    row_cells[0].text = var
    row_cells[1].text = purp
    row_cells[2].text = ex

add_row("SUPABASE_URL", "Database connection URL", "https://xyz.supabase.co")
add_row("SUPABASE_SERVICE_KEY", "Bypasses Row Level Security for backend operations", "eyJhbGciOiJIUzI1NiIsInR5c...")
add_row("OPENAI_API_KEY", "Primary AI orchestration & fallback gap analysis", "sk-proj-...")
add_row("GEMINI_API_KEY", "Used for high-context scraping limits", "AIzaSy...")

# --- SLACK INTEGRATION ---
doc.add_heading('4. Complete Slack Integration Guide', level=1)

doc.add_paragraph("To pipe Grayn AEO data into your team's Slack, you will need to create a custom Slack App via api.slack.com/apps and connect it to the Railway public URLs.")

doc.add_heading('Step 1: Create the Slack App', level=2)
doc.add_paragraph("1. Go to https://api.slack.com/apps and click 'Create New App'.", style='List Number')
doc.add_paragraph("2. Choose 'From scratch', name it 'Grayn AEO', and select your workspace.", style='List Number')

doc.add_heading('Step 2: Connect the Webhooks', level=2)
doc.add_paragraph("Grayn AEO endpoints output native Slack 'Block Kit' JSON payloads (`SlackPayload`). This means your middleware or serverless function just needs to GET data from Grayn and POST it directly to Slack without any formatting.")

doc.add_paragraph("You have 5 dedicated webhooks available:")
p = doc.add_paragraph()
p.add_run("• GET /v1/slack/daily?persona=exec: ").bold = True
p.add_run("Ideal for a daily cron job. Returns the overall Visibility % and weekly delta.\n")
p.add_run("• GET /v1/slack/content-ideas: ").bold = True
p.add_run("Returns the top 5 missing topic clusters.\n")
p.add_run("• GET /v1/slack/standing: ").bold = True
p.add_run("Returns the competitive Share of Voice leaderboard.\n")
p.add_run("• GET /v1/slack/competitor-sources: ").bold = True
p.add_run("Returns domains the AIs cite for your competitors.\n")

doc.add_heading('Step 3: Setting up the /aeo Slash Command', level=2)
doc.add_paragraph("We have a dedicated routing endpoint designed explicitly for Slack Slash Commands:")
p = doc.add_paragraph()
p.add_run("Target: POST /v1/slack/query").bold = True
doc.add_paragraph("How to configure it:", style='List Bullet')
doc.add_paragraph("1. In your Slack App settings, go to 'Slash Commands' and create `/aeo`.", style='List Bullet')
doc.add_paragraph("2. Slack will send a URL-encoded payload containing a `text` field whenever a user types `/aeo visibility`.", style='List Bullet')
doc.add_paragraph("3. Because Grayn AEO requires standard JSON and an Authorization Bearer token, you should use an API Gateway, Zapier, or a tiny serverless wrapper to intercept Slack's POST, convert the `text` field into the JSON format `{\"question\": \"[TEXT]\"}`, add your Grayn `Authorization: Bearer <API_KEY>` header, and forward the request to the Railway URL.", style='List Bullet')

doc.add_heading('Step 4: Authentication Middle-layer', level=2)
doc.add_paragraph("Crucially, every request hitting the Grayn AEO `/v1/` routes REQUIRES authentication. Because Slack does not natively allow you to inject custom HTTP Headers into its outgoing Slash Command webhooks, you MUST use a middleware layer (like Make.com, Zapier, or a lightweight AWS Lambda / Vercel Edge function) to catch the Slack payload, inject the `Authorization: Bearer <KEY>` header, and proxy it to the Railway backend.")

# Save the document
doc.save('Grayn_AEO_Comprehensive_Backend_Documentation.docx')
print("Document saved successfully.")
