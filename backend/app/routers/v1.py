"""
Grayn AEO — Public API Router (/v1)

All endpoints require a workspace Bearer key in the Authorization header.
Organised by domain matching the BRD's API surface (Appendix A).
"""

from fastapi import APIRouter, Depends, BackgroundTasks, Query, Response, HTTPException
import httpx
import uuid
import os
from fastapi.responses import JSONResponse
from supabase import Client
from app.database import get_supabase
from app.dependencies import get_current_workspace, verify_slack_api_key
from app.config import get_settings
from app.models.schemas import (
    WorkspaceOut,
    PromptCreate,
    PromptOut,
    PromptBulkCreate,
    PromptSeed,
    CompetitorCreate,
    CompetitorOut,
    RunTrigger,
    RunOut,
    RunStatus,
    RunLogItem,
    VisibilityScore,
    FullReport,
    CitationBreakdown,
    ClusterOut,
    ContentBrief,
    ContentDraft,
    Recommendation,
    AlertOut,
    DigestOut,
    SlackPayload,
    SlackQuery,
    LeaderboardEntry,
    Persona,
    DiscoverRequest,
    DiscoverResult,
    DiscoverApply,
    EngineType,
    WorkstreamCreate,
    WorkstreamOut,
    SlackScanTriggerRequest,
    SlackReportTriggerRequest
)
from app.services import scoring, tracking, discovery
import hashlib
import uuid

router = APIRouter(prefix="/v1", tags=["v1"])


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  Read / Metrics
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━


@router.get("/me", response_model=WorkspaceOut)
def get_me(workspace: dict = Depends(get_current_workspace)):
    """Return the current workspace profile."""
    return workspace


@router.get("/report", response_model=FullReport)
async def get_report(
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    """SC-10: Composite branded report combining all metrics."""
    return await scoring.build_full_report(db, workspace["id"])


@router.get("/trend")
def get_trend(
    weeks: int = Query(6, le=12),
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    """SC-13: Historical visibility trend data."""
    return scoring.compute_historical_trend(db, workspace["id"], weeks)


@router.get("/reports/visibility.pdf")
async def get_visibility_pdf(
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    from playwright.async_api import async_playwright
    import logging
    logger = logging.getLogger(__name__)
    
    url = f"http://localhost:3000/dashboard"
    
    try:
        async with async_playwright() as p:
            browser = await p.chromium.launch(headless=True)
            page = await browser.new_page()
            # Note: in real app, we'd pass auth token. Assuming local dev dashboard is accessible.
            await page.goto(url, wait_until="networkidle")
            pdf_bytes = await page.pdf(format="A4", print_background=True)
            await browser.close()
            
        return Response(content=pdf_bytes, media_type="application/pdf")
    except Exception as e:
        logger.error(f"PDF generation failed: {e}")
        return JSONResponse(status_code=500, content={"detail": "PDF generation failed"})


@router.get("/visibility", response_model=VisibilityScore)
def get_visibility(
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    """SC-01/02/03: Visibility %, delta, per-engine breakdown."""
    return scoring.compute_visibility(db, workspace["id"])


@router.get("/citations", response_model=list[CitationBreakdown])
def get_citations(
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    """SC-07: Citation source breakdown."""
    return scoring.compute_citations(db, workspace["id"])


@router.get("/competitors", response_model=list[CompetitorOut])
def list_competitors(
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    """WS-03: List tracked competitors."""
    result = (
        db.table("aeo_competitors")
        .select("*")
        .eq("workspace_id", workspace["id"])
        .execute()
    )
    return result.data or []


@router.get("/prompts", response_model=None)
def list_prompts(
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    """Get all prompts for the workspace."""
    result = (
        db.table("aeo_prompts")
        .select("id, prompt_text, topic_cluster, intent")
        .eq("workspace_id", workspace["id"])
        .eq("is_active", True)
        .order("created_at", desc=True)
        .execute()
    )
    return result.data or []


@router.get("/prompts/{prompt_id}/tracker", response_model=None)
async def query_tracker(
    prompt_id: str,
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    """Get in-depth query tracking data for a specific prompt."""
    return await scoring.get_query_data_tracker(db, workspace["id"], prompt_id)



@router.get("/competitors/sources")
def get_competitor_sources(
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    """SC-08: Domains AI cites for each competitor."""
    return scoring.compute_competitor_sources(db, workspace["id"])


@router.get("/clusters", response_model=list[ClusterOut])
def list_clusters(
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    """PC-03: List topic clusters."""
    result = (
        db.table("aeo_clusters")
        .select("*")
        .eq("workspace_id", workspace["id"])
        .order("opportunity_score", desc=True)
        .execute()
    )
    return result.data or []


@router.get("/clusters/{cluster_id}/brief", response_model=ContentBrief)
def get_cluster_brief(
    cluster_id: str,
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    """CB-01: Content brief for a topic cluster."""
    cluster = (
        db.table("aeo_clusters")
        .select("*")
        .eq("id", cluster_id)
        .eq("workspace_id", workspace["id"])
        .single()
        .execute()
        .data
    )

    # Get prompts in this cluster
    prompts = (
        db.table("aeo_prompts")
        .select("prompt_text")
        .eq("workspace_id", workspace["id"])
        .eq("topic_cluster", cluster["cluster_name"])
        .execute()
        .data
        or []
    )

    return ContentBrief(
        cluster_id=cluster["id"],
        cluster_name=cluster["cluster_name"],
        target_queries=[p["prompt_text"] for p in prompts],
        outline=[
            f"Introduction: What is {cluster['cluster_name']}?",
            "Key features and comparison of top solutions",
            "Detailed analysis with pros and cons",
            "Expert recommendations and use cases",
            "FAQ section addressing common questions",
        ],
        sources_to_win=[
            "Create comprehensive, expert-level content",
            "Include specific data points and benchmarks",
            "Add structured FAQ section for AI extraction",
            "Use schema markup for enhanced visibility",
        ],
        refill_action=cluster.get("refill_action", "write-new"),
    )

from app.services.content_analyzer import analyze_content_gaps
from pydantic import BaseModel

class ContentGapResponse(BaseModel):
    brief_markdown: str
    urls_analyzed: list[str]

@router.get("/content/gaps", response_model=ContentGapResponse)
async def get_content_gaps(
    prompt_id: str,
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    """PC-06: Scrape top cited URLs and generate organic content gap brief."""
    ws_id = workspace["id"]

    # 1. Get the prompt text
    prompt_res = db.table("aeo_prompts").select("prompt_text").eq("id", prompt_id).eq("workspace_id", ws_id).execute()
    if not prompt_res.data:
        raise HTTPException(status_code=404, detail="Prompt not found")
    prompt_text = prompt_res.data[0]["prompt_text"]

    # 2. Get top 5 unique cited URLs for this prompt across all recent runs
    # Find all run IDs for this prompt
    runs = db.table("aeo_runs").select("id").eq("prompt_id", prompt_id).eq("workspace_id", ws_id).execute().data or []
    if not runs:
        return ContentGapResponse(brief_markdown="No tracking data exists for this prompt yet. Please trigger a tracking run first.", urls_analyzed=[])
    
    run_ids = [r["id"] for r in runs]
    
    # Supabase Python client doesn't support generic 'in_' well sometimes, but we can try or fetch all and filter
    citations = db.table("aeo_citations").select("url").eq("workspace_id", ws_id).in_("run_id", run_ids).execute().data or []
    
    unique_urls = list(set([c["url"] for c in citations if c.get("url")]))
    # Filter out empty or obviously bad URLs, take top 15 for diverse data
    urls_to_scrape = unique_urls[:15]

    if not urls_to_scrape:
        return ContentGapResponse(brief_markdown="No competitor citations found for this prompt. You might already own this space!", urls_analyzed=[])

    # 3. Analyze content gaps
    brief = await analyze_content_gaps(urls=urls_to_scrape, prompt_text=prompt_text, brand_name=workspace.get("brand_name", ""))

    return ContentGapResponse(
        brief_markdown=brief,
        urls_analyzed=urls_to_scrape
    )


@router.get("/recommendations", response_model=list[Recommendation])
def get_recommendations(
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    """PC-05: Prioritized content gap recommendations."""
    clusters = (
        db.table("aeo_clusters")
        .select("*")
        .eq("workspace_id", workspace["id"])
        .order("opportunity_score", desc=True)
        .execute()
        .data
        or []
    )

    recommendations = []
    for c in clusters:
        prompts = (
            db.table("aeo_prompts")
            .select("prompt_text")
            .eq("workspace_id", workspace["id"])
            .eq("topic_cluster", c["cluster_name"])
            .limit(5)
            .execute()
            .data
            or []
        )
        recommendations.append(
            Recommendation(
                cluster_id=c["id"],
                cluster_name=c["cluster_name"],
                action=c.get("refill_action", "write-new"),
                impact_score=c.get("opportunity_score", 0) or 0,
                target_queries=[p["prompt_text"] for p in prompts],
            )
        )

    return recommendations


@router.get("/runs", response_model=list[RunOut])
def list_runs(
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
    limit: int = Query(default=50, le=200),
):
    """List recent tracking runs."""
    result = (
        db.table("aeo_runs")
        .select("*")
        .eq("workspace_id", workspace["id"])
        .order("created_at", desc=True)
        .limit(limit)
        .execute()
    )
    return result.data or []


@router.get("/prompts", response_model=list[PromptOut])
def list_prompts(
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    """QP-01: List all prompts."""
    result = (
        db.table("aeo_prompts")
        .select("*")
        .eq("workspace_id", workspace["id"])
        .order("created_at", desc=True)
        .execute()
    )
    return result.data or []


@router.get("/alerts", response_model=list[AlertOut])
def list_alerts(
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    """AL-01: Threshold alerts."""
    vis = scoring.compute_visibility(db, workspace["id"])
    alerts = []
    if vis.week_over_week_delta is not None and vis.week_over_week_delta < -5:
        alerts.append(
            AlertOut(
                type="visibility_drop",
                message=f"Visibility dropped by {abs(vis.week_over_week_delta)}% this week",
                severity="warning",
            )
        )
    return alerts


@router.get("/digest", response_model=DigestOut | None)
def get_latest_digest(
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    """AL-02: Latest weekly digest."""
    result = (
        db.table("aeo_digests")
        .select("*")
        .eq("workspace_id", workspace["id"])
        .order("created_at", desc=True)
        .limit(1)
        .maybe_single()
        .execute()
    )
    return result.data if result else None


@router.get("/digests", response_model=list[DigestOut])
def list_digests(
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    """List all digests."""
    result = (
        db.table("aeo_digests")
        .select("*")
        .eq("workspace_id", workspace["id"])
        .order("created_at", desc=True)
        .execute()
    )
    return result.data or []


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  Write / Actions
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━


@router.post("/prompts", response_model=PromptOut)
def create_prompt(
    body: PromptCreate,
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    """QP-01: Create a single prompt."""
    data = {
        "workspace_id": workspace["id"],
        "prompt_text": body.prompt_text.strip().lower(),
        "intent": body.intent.value if body.intent else None,
        "persona": body.persona.value if body.persona else None,
        "topic_cluster": body.topic_cluster,
    }
    result = db.table("aeo_prompts").insert(data).execute()
    return result.data[0]


@router.post("/prompts/bulk", response_model=list[PromptOut])
def create_prompts_bulk(
    body: PromptBulkCreate,
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    """QP-02: Bulk prompt creation."""
    rows = [
        {
            "workspace_id": workspace["id"],
            "prompt_text": p.prompt_text.strip().lower(),
            "intent": p.intent.value if p.intent else None,
            "persona": p.persona.value if p.persona else None,
            "topic_cluster": p.topic_cluster,
        }
        for p in body.prompts
    ]
    result = db.table("aeo_prompts").insert(rows).execute()
    return result.data


@router.post("/competitors", response_model=CompetitorOut)
def add_competitor(
    body: CompetitorCreate,
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    """WS-03: Add a tracked competitor."""
    data = {
        "workspace_id": workspace["id"],
        "brand_name": body.brand_name,
        "domain": body.domain,
        "aliases": body.aliases,
    }
    result = db.table("aeo_competitors").insert(data).execute()
    return result.data[0]


@router.post("/runs/trigger")
async def trigger_runs(
    body: RunTrigger,
    background_tasks: BackgroundTasks,
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    """TR-08: Trigger tracking runs (background)."""

    async def _run():
        await tracking.trigger_batch_run(
            db, workspace, body.engines, body.prompt_ids
        )

    background_tasks.add_task(_run)
    return {"status": "triggered", "message": "Batch run started in background"}


@router.get("/runs/status", response_model=RunStatus)
def get_run_status(
    engines_count: int = Query(5, description="Number of engines to track"),
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    """TR-09: Live progress status for the current week's batch run."""
    from app.services.tracking import _current_iso_week
    iso_week = _current_iso_week()
    workspace_id = workspace["id"]

    # Calculate total expected runs
    active_prompts = (
        db.table("aeo_prompts")
        .select("id", count="exact")
        .eq("workspace_id", workspace_id)
        .eq("is_active", True)
        .execute()
    )
    total_prompts = active_prompts.count or 0

    total_expected = total_prompts * engines_count

    # Get completed runs for this week
    completed_runs = (
        db.table("aeo_runs")
        .select("id", count="exact")
        .eq("workspace_id", workspace_id)
        .eq("iso_week", iso_week)
        .execute()
    )
    completed_count = completed_runs.count or 0

    # Get latest logs
    latest_runs = (
        db.table("aeo_runs")
        .select("engine, status, created_at, aeo_prompts(prompt_text)")
        .eq("workspace_id", workspace_id)
        .eq("iso_week", iso_week)
        .order("created_at", desc=True)
        .limit(50)
        .execute()
    )

    logs = []
    for r in (latest_runs.data or []):
        prompt_text = r.get("aeo_prompts", {}).get("prompt_text", "Unknown query")
        logs.append(RunLogItem(
            engine=r["engine"],
            prompt_text=prompt_text,
            status=r["status"],
            created_at=r["created_at"],
        ))

    progress_pct = 0.0
    if total_expected > 0:
        progress_pct = round(min(100.0, (completed_count / total_expected) * 100), 2)
    elif total_expected == 0 and completed_count > 0:
        progress_pct = 100.0

    return RunStatus(
        progress_pct=progress_pct,
        completed=completed_count,
        total=total_expected,
        is_running=(progress_pct < 100.0 and progress_pct > 0.0),
        latest_logs=logs
    )


@router.post("/clusters/{cluster_id}/draft", response_model=ContentDraft)
def generate_draft(
    cluster_id: str,
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    """CB-02: Generate an AEO-ready draft for a topic cluster."""
    cluster = (
        db.table("aeo_clusters")
        .select("*")
        .eq("id", cluster_id)
        .eq("workspace_id", workspace["id"])
        .single()
        .execute()
        .data
    )

    prompts = (
        db.table("aeo_prompts")
        .select("prompt_text")
        .eq("workspace_id", workspace["id"])
        .eq("topic_cluster", cluster["cluster_name"])
        .execute()
        .data
        or []
    )

    questions = [p["prompt_text"] for p in prompts]
    brand = workspace.get("brand_name", "Your Brand")

    # Template-based draft (CB-04 will swap for a real LLM writer)
    body_sections = []
    body_sections.append(
        f"# {cluster['cluster_name']}: The Complete Guide\n\n"
        f"As businesses look for the best solutions in {cluster['cluster_name']}, "
        f"{brand} stands out with its comprehensive approach.\n"
    )

    if questions:
        body_sections.append("## Frequently Asked Questions\n")
        for q in questions[:5]:
            body_sections.append(
                f"### {q}\n\n"
                f"[Expert answer addressing this question with specific data "
                f"points, comparisons, and actionable recommendations.]\n"
            )

    body_sections.append(
        f"\n## Why {brand}?\n\n"
        f"{brand} provides industry-leading solutions backed by data-driven "
        f"insights and customer success stories.\n\n"
        f"## Next Steps\n\n"
        f"Ready to get started? Contact our team for a personalised demo.\n"
    )

    return ContentDraft(
        cluster_id=cluster["id"],
        cluster_name=cluster["cluster_name"],
        title=f"{cluster['cluster_name']}: The Complete Guide",
        body="\n".join(body_sections),
    )


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  Workstreams
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━


@router.get("/workstreams", response_model=list[WorkstreamOut])
def list_workstreams(
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    ws_id = workspace["id"]
    res = db.table("aeo_workstreams").select("*").eq("workspace_id", ws_id).execute()
    return res.data


@router.post("/workstreams", response_model=WorkstreamOut)
def create_workstream(
    payload: WorkstreamCreate,
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    ws_id = workspace["id"]
    data = {
        "workspace_id": ws_id,
        "name": payload.name,
        "topics": payload.topics,
        "attribute_filters": payload.attribute_filters,
        "target_visibility": payload.target_visibility,
    }
    res = db.table("aeo_workstreams").insert(data).execute()
    if not res.data:
        raise HTTPException(status_code=500, detail="Failed to create workstream")
    return res.data[0]


@router.get("/workstreams/{workstream_id}", response_model=WorkstreamOut)
def get_workstream(
    workstream_id: str,
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    ws_id = workspace["id"]
    res = db.table("aeo_workstreams").select("*").eq("id", workstream_id).eq("workspace_id", ws_id).execute()
    if not res.data:
        raise HTTPException(status_code=404, detail="Workstream not found")
    return res.data[0]


@router.patch("/workstreams/{workstream_id}", response_model=WorkstreamOut)
def update_workstream(
    workstream_id: str,
    payload: WorkstreamCreate,
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    ws_id = workspace["id"]
    data = {
        "name": payload.name,
        "topics": payload.topics,
        "attribute_filters": payload.attribute_filters,
        "target_visibility": payload.target_visibility,
    }
    res = db.table("aeo_workstreams").update(data).eq("id", workstream_id).eq("workspace_id", ws_id).execute()
    if not res.data:
        raise HTTPException(status_code=404, detail="Workstream not found")
    return res.data[0]


@router.delete("/workstreams/{workstream_id}")
def delete_workstream(
    workstream_id: str,
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    ws_id = workspace["id"]
    db.table("aeo_workstreams").delete().eq("id", workstream_id).eq("workspace_id", ws_id).execute()
    return {"status": "deleted"}


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  Slack Payloads
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━


@router.get("/slack/daily", response_model=SlackPayload)
def slack_daily_update(
    persona: Persona = Query(default=Persona.EXEC),
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    """SL-01: Daily/weekly visibility update as Block Kit."""
    vis = scoring.compute_visibility(db, workspace["id"])
    sov, leaderboard = scoring.compute_share_of_voice(db, workspace["id"])

    brand_name = workspace.get("brand_name", "Your Brand")
    delta_str = ""
    if vis.week_over_week_delta is not None:
        sign = "+" if vis.week_over_week_delta >= 0 else ""
        delta_str = f" ({sign}{vis.week_over_week_delta}% vs last week)"

    if persona == Persona.EXEC:
        blocks = [
            {
                "type": "header",
                "text": {
                    "type": "plain_text",
                    "text": f"📊 {brand_name} — AI Visibility Update",
                },
            },
            {
                "type": "section",
                "text": {
                    "type": "mrkdwn",
                    "text": (
                        f"*Visibility:* {vis.visibility_pct}%{delta_str}\n"
                        f"*Week:* {vis.iso_week}"
                    ),
                },
            },
        ]
        if leaderboard:
            lb_text = "\n".join(
                f"{e.rank}. {e.brand_name} — {e.share_pct}%"
                for e in leaderboard[:5]
            )
            blocks.append({
                "type": "section",
                "text": {"type": "mrkdwn", "text": f"*Leaderboard:*\n{lb_text}"},
            })
    else:
        # SEO persona — more detail
        engine_text = "\n".join(
            f"• {eng}: {pct}%" for eng, pct in vis.per_engine.items()
        )
        blocks = [
            {
                "type": "header",
                "text": {
                    "type": "plain_text",
                    "text": f"🔍 {brand_name} — Detailed AI Visibility",
                },
            },
            {
                "type": "section",
                "text": {
                    "type": "mrkdwn",
                    "text": (
                        f"*Overall Visibility:* {vis.visibility_pct}%{delta_str}\n"
                        f"*Week:* {vis.iso_week}\n\n"
                        f"*Per Engine:*\n{engine_text}"
                    ),
                },
            },
        ]
        if leaderboard:
            lb_text = "\n".join(
                f"{e.rank}. {e.brand_name} — {e.share_pct}% "
                f"(avg pos: {e.avg_position or 'N/A'})"
                for e in leaderboard
            )
            blocks.append({
                "type": "section",
                "text": {"type": "mrkdwn", "text": f"*Full Leaderboard:*\n{lb_text}"},
            })

    return SlackPayload(
        blocks=blocks,
        text=f"{brand_name} AI Visibility: {vis.visibility_pct}%{delta_str}",
    )


@router.get("/slack/content-ideas", response_model=SlackPayload)
def slack_content_ideas(
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    """SL-02: Top content gaps as a Slack message."""
    clusters = (
        db.table("aeo_clusters")
        .select("*")
        .eq("workspace_id", workspace["id"])
        .order("opportunity_score", desc=True)
        .limit(5)
        .execute()
        .data
        or []
    )

    items = "\n".join(
        f"• *{c['cluster_name']}* — Score: {c.get('opportunity_score', 0)} | Action: {c.get('refill_action', 'write-new')}"
        for c in clusters
    ) or "No content gaps identified yet. Run tracking first!"

    blocks = [
        {
            "type": "header",
            "text": {"type": "plain_text", "text": "💡 Content Ideas"},
        },
        {
            "type": "section",
            "text": {"type": "mrkdwn", "text": items},
        },
    ]

    return SlackPayload(blocks=blocks, text="Top content ideas from Grayn AEO")


@router.get("/slack/standing", response_model=SlackPayload)
def slack_standing(
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    """SL-03: Where the brand ranks vs. competitors."""
    _, leaderboard = scoring.compute_share_of_voice(db, workspace["id"])
    brand_name = workspace.get("brand_name", "Your Brand")

    lb_text = "\n".join(
        f"{'👑 ' if e.brand_name == brand_name else ''}"
        f"{e.rank}. {e.brand_name} — {e.share_pct}%"
        for e in leaderboard
    ) or "No data yet. Run tracking first!"

    blocks = [
        {
            "type": "header",
            "text": {"type": "plain_text", "text": f"🏆 {brand_name} — Competitive Standing"},
        },
        {
            "type": "section",
            "text": {"type": "mrkdwn", "text": lb_text},
        },
    ]

    return SlackPayload(blocks=blocks, text=f"{brand_name} competitive standing")


@router.get("/slack/competitor-sources", response_model=SlackPayload)
def slack_competitor_sources(
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    """SL-04: Where rivals get cited."""
    comp_sources = scoring.compute_competitor_sources(db, workspace["id"])

    sections = []
    for comp, sources in comp_sources.items():
        domains = set(s["domain"] for s in sources)
        sections.append(f"*{comp}:* {', '.join(list(domains)[:5])}")

    text = "\n".join(sections) or "No competitor sources identified yet."

    blocks = [
        {
            "type": "header",
            "text": {"type": "plain_text", "text": "🔗 Competitor Sources"},
        },
        {
            "type": "section",
            "text": {"type": "mrkdwn", "text": text},
        },
    ]

    return SlackPayload(blocks=blocks, text="Competitor source intelligence")


@router.post("/slack/query", response_model=SlackPayload)
def slack_query_router(
    body: SlackQuery,
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase),
):
    """SL-05: Free-text /aeo query router."""
    question = body.question.lower()
    brand_name = workspace.get("brand_name", "Your Brand")

    # Simple intent routing
    if any(word in question for word in ["visibility", "visible", "mentioned"]):
        vis = scoring.compute_visibility(db, workspace["id"])
        answer = f"Your current visibility is *{vis.visibility_pct}%* (week {vis.iso_week})."
    elif any(word in question for word in ["rank", "position", "standing", "leaderboard"]):
        _, lb = scoring.compute_share_of_voice(db, workspace["id"])
        brand_entry = next(
            (e for e in lb if e.brand_name == brand_name), None
        )
        if brand_entry:
            answer = f"You're ranked *#{brand_entry.rank}* with *{brand_entry.share_pct}%* share of voice."
        else:
            answer = "No ranking data yet. Run tracking first!"
    elif any(word in question for word in ["competitor", "rival", "who"]):
        comps = (
            db.table("aeo_competitors")
            .select("brand_name")
            .eq("workspace_id", workspace["id"])
            .execute()
            .data
            or []
        )
        names = ", ".join(c["brand_name"] for c in comps)
        answer = f"Your tracked competitors: {names}" if names else "No competitors tracked yet."
    elif any(word in question for word in ["content", "write", "gap", "idea"]):
        clusters = (
            db.table("aeo_clusters")
            .select("cluster_name, opportunity_score")
            .eq("workspace_id", workspace["id"])
            .order("opportunity_score", desc=True)
            .limit(3)
            .execute()
            .data
            or []
        )
        if clusters:
            items = ", ".join(
                f"*{c['cluster_name']}* (score: {c.get('opportunity_score', 0)})"
                for c in clusters
            )
            answer = f"Top content opportunities: {items}"
        else:
            answer = "No content gaps identified yet."
    else:
        answer = (
            f"I can help you with: visibility, ranking, competitors, and content gaps. "
            f"Try asking something like 'how visible are we?' or 'who are our competitors?'"
        )

    blocks = [
        {
            "type": "section",
            "text": {"type": "mrkdwn", "text": answer},
        },
    ]

    return SlackPayload(blocks=blocks, text=answer)


# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  Discovery / Onboarding
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━


@router.post("/discover", response_model=DiscoverResult)
async def run_discovery(body: DiscoverRequest):
    """Fetch URL and extract brand context automatically."""
    try:
        parsed = await discovery.run_discovery(body.url, num_queries=body.num_queries)
        return DiscoverResult(
            brand_name=parsed.get("brand_name"),
            suggested_competitors=parsed.get("suggested_competitors", []),
            suggested_queries=parsed.get("suggested_queries", []),
            themes=parsed.get("themes", [])
        )
    except Exception as e:
        from fastapi import HTTPException
        raise HTTPException(status_code=400, detail=str(e))


@router.post("/workspaces/onboard")
async def onboard_workspace(
    body: DiscoverApply,
    background_tasks: BackgroundTasks,
    workspace: dict = Depends(get_current_workspace),
    db: Client = Depends(get_supabase)
):
    """Create a workspace and trigger initial tracking runs."""
    # Since we use Depends(get_current_workspace), the user must have *some* token.
    # Actually, onboarding might be public or require a generic token.
    # In OpenLens style, they enter a URL directly. 
    # Let's mint a new workspace regardless.
    
    brand_name = body.competitors[0].brand_name if body.competitors else "Unknown Brand"
    domain = body.competitors[0].domain if body.competitors else ""
    
    # We will just overwrite the current workspace for simplicity in this demo,
    # or better, clear existing data for this workspace and insert the new stuff.
    ws_id = workspace["id"]
    
    # Update Workspace
    update_data = {
        "brand_name": brand_name,
        "domain": domain,
        "aliases": [brand_name]
    }
    if body.target_location:
        update_data["target_location"] = body.target_location
        
    db.table("workspaces").update(update_data).eq("id", ws_id).execute()

    # Clear old data (for clean onboarding)
    db.table("aeo_competitors").delete().eq("workspace_id", ws_id).execute()
    db.table("aeo_clusters").delete().eq("workspace_id", ws_id).execute()
    db.table("aeo_prompts").delete().eq("workspace_id", ws_id).execute()

    # Insert new Competitors
    if body.competitors:
        # Skip the first one as it is the brand itself (usually)
        comps = body.competitors[1:]
        if comps:
            db.table("aeo_competitors").insert([
                {"workspace_id": ws_id, "brand_name": c.brand_name, "domain": c.domain, "aliases": c.aliases}
                for c in comps
            ]).execute()
            
    # Insert Themes as Clusters
    themes = list(set(["General Industry"] * len(body.queries))) # Mocking themes if needed, or we can just use "General"
    db.table("aeo_clusters").insert({
        "workspace_id": ws_id,
        "cluster_name": "Discovery Queries"
    }).execute()

    # Insert Prompts
    if body.queries:
        db.table("aeo_prompts").insert([
            {
                "workspace_id": ws_id, 
                "prompt_text": q.text, 
                "topic_cluster": "Discovery Queries",
                "attributes": q.attributes
            }
            for q in body.queries
        ]).execute()
        
    # Get all prompt IDs
    prompt_ids = [p["id"] for p in db.table("aeo_prompts").select("id").eq("workspace_id", ws_id).execute().data]

    # Trigger runs across selected engines
    engines = body.engines if body.engines else [EngineType.OPENAI, EngineType.GEMINI, EngineType.GOOGLE_AI, EngineType.PERPLEXITY, EngineType.GROK]
    
    async def _run_all():
        await tracking.trigger_batch_run(db, workspace, engines, prompt_ids)
            
    background_tasks.add_task(_run_all)
    
    return {"status": "onboarded", "workspace_id": ws_id, "message": "Tracking started in background"}



# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
#  Slack External Handlers (Lovable App Contract)
# ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

@router.post("/scans/trigger", status_code=202)
async def trigger_slack_scan(
    body: SlackScanTriggerRequest,
    background_tasks: BackgroundTasks,
    is_valid: bool = Depends(verify_slack_api_key),
    db: Client = Depends(get_supabase)
):
    scan_id = str(uuid.uuid4())
    
    async def _run_scan_bg():
        payload = {
            "flow_key": "trigger_new_ai_scan",
            "workspace_id": body.workspace_id,
            "channel": body.callback.get("channel"),
            "thread_ts": body.callback.get("thread_ts"),
            "status": "ok",
            "payload": {"scan_id": scan_id}
        }
        try:
            ws_res = db.table("workspaces").select("*").eq("id", body.workspace_id).execute()
            if not ws_res.data:
                raise ValueError("Workspace not found")
            ws = ws_res.data[0]
            
            prompt_res = db.table("aeo_prompts").select("id").eq("workspace_id", body.workspace_id).execute()
            prompt_ids = [p["id"] for p in prompt_res.data]
            
            if prompt_ids:
                engines = [EngineType.OPENAI, EngineType.GEMINI, EngineType.GOOGLE_AI, EngineType.PERPLEXITY]
                await tracking.trigger_batch_run(db, ws, engines, prompt_ids)
                
        except Exception as e:
            payload["status"] = "error"
            payload["payload"] = {"error": str(e)}
            
        # Emit callback
        settings = get_settings()
        async with httpx.AsyncClient() as client:
            await client.post(
                f"{settings.SUPABASE_URL}/functions/v1/slack-flow-callback",
                headers={"Authorization": f"Bearer {settings.GRAYN_AEO_API_KEY}"},
                json=payload
            )

    background_tasks.add_task(_run_scan_bg)
    return {"scan_id": scan_id, "status": "queued"}


@router.post("/reports/full-docx", status_code=202)
async def trigger_slack_report(
    body: SlackReportTriggerRequest,
    background_tasks: BackgroundTasks,
    is_valid: bool = Depends(verify_slack_api_key),
    db: Client = Depends(get_supabase)
):
    report_id = str(uuid.uuid4())
    
    async def _run_report_bg():
        payload = {
            "flow_key": "generate_full_report",
            "workspace_id": body.workspace_id,
            "channel": body.callback.get("channel"),
            "thread_ts": body.callback.get("thread_ts"),
            "status": "ok",
            "payload": {}
        }
        
        try:
            from docx import Document
            import tempfile
            
            workspace_id = body.workspace_id
            
            # Fetch data
            report = await scoring.build_full_report(db, workspace_id)
            trend = scoring.compute_historical_trend(db, workspace_id)
            prompts = db.table("aeo_prompts").select("*").eq("workspace_id", workspace_id).execute().data or []
            clusters = db.table("aeo_clusters").select("*").eq("workspace_id", workspace_id).execute().data or []
            runs = db.table("aeo_runs").select("id, engine, prompt_id, created_at").eq("workspace_id", workspace_id).order("created_at", desc=True).limit(50).execute().data or []
            
            mentions_data = []
            for run in runs:
                m = db.table("aeo_mentions").select("brand_name, sentiment, position, is_target_brand").eq("run_id", run["id"]).execute().data or []
                mentions_data.append({"run": run, "mentions": m})
                
            raw_payload = {
                "report": report.model_dump(mode='json'),
                "trend": trend,
                "prompts": prompts,
                "clusters": clusters,
                "runs_and_mentions": mentions_data
            }
            
            # Create Docx
            document = Document()
            document.add_heading(f'AEO Detailed Report - {body.brand_name}', 0)
            
            # 1. Overview
            document.add_heading('1. Overview Dashboard', level=1)
            rep = raw_payload["report"]
            document.add_paragraph(f"Overall Visibility Score: {rep['visibility']['visibility_pct']}%")
            document.add_paragraph(f"Week-over-week delta: {rep['visibility'].get('week_over_week_delta', 0)}%")
            for engine, pct in rep['visibility']['per_engine'].items():
                document.add_paragraph(f"- {engine}: {pct}%")
            
            # 2. Topic Clusters
            document.add_heading('2. Topic Clusters', level=1)
            if clusters:
                for c in clusters:
                    document.add_paragraph(f"Cluster: {c.get('cluster_name')} | Search Volume: {c.get('search_volume', 'N/A')} | Opportunity Score: {c.get('opportunity_score', 'N/A')}")
            else:
                document.add_paragraph("No topic clusters found.")
                
            # 3. Content Gaps
            document.add_heading('3. Content Gaps Studio', level=1)
            if prompts:
                for p in prompts[:5]:
                    document.add_paragraph(f"- The query '{p['prompt_text']}' lacks strong informational content linking back to {body.brand_name}.")
            else:
                document.add_paragraph("No prompts found to analyze gaps.")
                
            # 4. Query Manager
            document.add_heading('4. Query Manager', level=1)
            for p in prompts:
                document.add_paragraph(f"Query: {p['prompt_text']} | Intent: {p.get('intent', 'unknown')} | Active: {p.get('is_active', True)}")
                
            # 5. Competitor Analysis
            document.add_heading('5. Competitor Analysis', level=1)
            if rep.get("leaderboard"):
                for c in rep["leaderboard"]:
                    document.add_paragraph(f"Rank #{c['rank']}: {c['brand_name']} - Share: {c['share_pct']}% - Mentions: {c['mention_count']}")
            else:
                document.add_paragraph("No competitors found.")
                
            # Save to tmp
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                doc_path = tmp.name
            document.save(doc_path)
            
            # Upload to Supabase Storage
            storage_path = f"reports/{workspace_id}/{report_id}.docx"
            with open(doc_path, 'rb') as f:
                res = db.storage.from_("grayn-aeo-artifacts").upload(storage_path, f, file_options={"content-type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"})
            
            os.remove(doc_path)
            
            payload["payload"] = {
                "storage_bucket": "grayn-aeo-artifacts",
                "storage_path": storage_path,
                "filename": f"{body.brand_name}_AEO_Report.docx",
                "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            }
            
        except Exception as e:
            payload["status"] = "error"
            payload["payload"] = {"error": str(e)}
            
        # Emit callback
        settings = get_settings()
        async with httpx.AsyncClient() as client:
            await client.post(
                f"{settings.SUPABASE_URL}/functions/v1/slack-flow-callback",
                headers={"Authorization": f"Bearer {settings.GRAYN_AEO_API_KEY}"},
                json=payload
            )

    background_tasks.add_task(_run_report_bg)
    return {"report_id": report_id, "status": "queued"}
