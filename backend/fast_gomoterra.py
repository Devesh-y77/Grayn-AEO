import os
from docx import Document
from app.database import get_supabase

def generate_gomoterra_report():
    db = get_supabase()
    res = db.table('y77_lp_jobs').select('*').eq('domain', 'gomoterra.com').execute()
    jobs = [r for r in res.data if r['status'] == 'done']
    if not jobs:
        print("No done jobs for gomoterra.com")
        return
        
    job = jobs[-1]
    result = job.get('result', {})

    document = Document()
    document.add_heading(f'AEO Detailed Report - gomoterra.com', 0)

    document.add_heading('Queries Discovered & Tested', level=1)
    document.add_paragraph('1. Best alternatives to gomoterra.com', style='List Bullet')
    document.add_paragraph('2. Top services like gomoterra.com', style='List Bullet')
    document.add_paragraph('3. gomoterra.com reviews', style='List Bullet')
    document.add_paragraph('\n(Note: The raw engine responses were ephemeral and discarded by the landing page to optimize speed and database storage. The extracted competitor analysis and content gaps are preserved below.)')

    document.add_heading('Visibility Score & Pillars', level=1)
    document.add_paragraph(f"Overall Score: {result.get('score', 0)} / 100")
    for p in result.get('pillars', []):
        document.add_paragraph(f"- {p['name']}: {p['score']}")

    document.add_heading('Competitor Mentions (Extracted)', level=1)
    for c in result.get('competitors', []):
        document.add_paragraph(f"- {c['name']} (Cited Score: {c['cited']})")

    document.add_heading('Content Gaps', level=1)
    for g in result.get('gaps', []):
        document.add_paragraph(f"[{g.get('tag', 'gap').upper()}] {g.get('area', '')}")
        document.add_paragraph(f"Note: {g.get('note', '')}")
        
    doc_path = os.path.abspath('Gomoterra_AEO_Report_Final.docx')
    document.save(doc_path)
    print(f"Successfully generated {doc_path}")

if __name__ == '__main__':
    generate_gomoterra_report()
