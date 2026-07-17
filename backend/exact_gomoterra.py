import os
from docx import Document
from app.database import get_supabase

def generate_gomoterra_report():
    db = get_supabase()
    # Fetch ordered by created_at desc so we get the ACTUAL latest
    res = db.table('y77_lp_jobs').select('*').eq('domain', 'gomoterra.com').order('created_at', desc=True).execute()
    jobs = [r for r in res.data if r['status'] == 'done']
    if not jobs:
        print("No done jobs for gomoterra.com")
        return
        
    job = jobs[0] # The latest one
    result = job.get('result', {})

    document = Document()
    document.add_heading(f'AEO Detailed Report - gomoterra.com', 0)

    document.add_paragraph('**IMPORTANT NOTE:** This document contains the exact data displayed on your latest website run. The raw AI engine text responses are not included because the website deletes them to save space immediately after scoring. Only this final analysis is saved.')

    document.add_heading('Visibility Score & Pillars', level=1)
    document.add_paragraph(f"Overall Score: {result.get('score', 0)} / 100")
    for p in result.get('pillars', []):
        document.add_paragraph(f"- {p['name']}: {p['score']}")

    document.add_heading('Competitor Mentions (Extracted)', level=1)
    for c in result.get('competitors', []):
        document.add_paragraph(f"- {c['name']} (Cited Score: {c['cited']})")

    document.add_heading('Content Gaps', level=1)
    for g in result.get('gaps', []):
        document.add_paragraph(f"[{g.get('tag', 'gap').upper()}] {g.get('area', '')}")
        document.add_paragraph(f"Note: {g.get('note', '')}")
        
    doc_path = os.path.abspath('Gomoterra_Exact_Website_Report.docx')
    document.save(doc_path)
    print(f"Successfully generated {doc_path}")

if __name__ == '__main__':
    generate_gomoterra_report()
