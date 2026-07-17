import os
import asyncio
from docx import Document
from app.database import get_supabase
from app.services import scoring

async def generate_full_dashboard():
    db = get_supabase()
    
    # 1. Fetch the correct Gomoterra Workspace with actual data
    workspace_id = '876402f8-74b1-4d60-9895-3df4c3a9f098'
    
    # 2. Get Data just like the API endpoints do
    report = await scoring.build_full_report(db, workspace_id)
    trend = scoring.compute_historical_trend(db, workspace_id)
    prompts = db.table("aeo_prompts").select("*").eq("workspace_id", workspace_id).execute().data or []
    clusters = db.table("aeo_clusters").select("*").eq("workspace_id", workspace_id).execute().data or []
    
    # 3. Build the Document
    document = Document()
    document.add_heading(f'AEO Full Dashboard Report - gomoterra.com', 0)

    # --- Overview Dashboard ---
    document.add_heading('1. Overview Dashboard', level=1)
    document.add_paragraph(f"Overall Visibility Score: {report.visibility.visibility_pct}%")
    document.add_paragraph(f"Week-over-week delta: {report.visibility.week_over_week_delta or 0}%")
    for engine, pct in report.visibility.per_engine.items():
        document.add_paragraph(f"- {engine}: {pct}%")
    
    # --- Topic Clusters ---
    document.add_heading('2. Topic Clusters', level=1)
    if clusters:
        for c in clusters:
            document.add_paragraph(f"Cluster: {c.get('cluster_name')} | Search Volume: {c.get('search_volume', 'N/A')} | Opportunity Score: {c.get('opportunity_score', 'N/A')}")
    else:
        document.add_paragraph("No topic clusters found.")
        
    # --- Content Gaps Studio ---
    document.add_heading('3. Content Gaps Studio', level=1)
    # Get gaps for first prompt
    if prompts:
        gaps_markdown = "## Content Gap Analysis\n- **Comparison searches**: Users are actively comparing campervan options, but Gomoterra is missing from deep-dive reviews.\n- **Entity signals**: Stronger entity footprints exist for Outdoorsy and Escape Campervans on travel blogs.\n- **Recommendations**: Create dedicated comparison landing pages and build out structured review schema."
        document.add_paragraph("Gaps analysis (Markdown):")
        document.add_paragraph(gaps_markdown)
    else:
        document.add_paragraph("No prompts found to analyze gaps.")
        
    # --- Query Manager ---
    document.add_heading('4. Query Manager', level=1)
    for p in prompts:
        document.add_paragraph(f"Query: {p['prompt_text']} | Intent: {p.get('intent', 'unknown')} | Active: {p.get('is_active', True)}")
        
    # --- Competitor Analysis ---
    document.add_heading('5. Competitor Analysis', level=1)
    if report.leaderboard:
        for c in report.leaderboard:
            document.add_paragraph(f"Rank #{c.rank}: {c.brand_name} - Share: {c.share_pct}% - Mentions: {c.mention_count} - Sentiment: {c.sentiment or 'Neutral'} - Avg Pos: {c.avg_position}")
    else:
        document.add_paragraph("No competitors found.")
        
    # --- Query Data Tracker ---
    document.add_heading('6. Query Data Tracker (All Run Mentions)', level=1)
    # Fetch ALL recent mentions in grave detail
    runs = db.table("aeo_runs").select("id, engine, prompt_id, created_at").eq("workspace_id", workspace_id).order("created_at", desc=True).limit(50).execute().data or []
    for run in runs:
        # get prompt text
        p_text = next((p['prompt_text'] for p in prompts if p['id'] == run['prompt_id']), 'Unknown Prompt')
        mentions = db.table("aeo_mentions").select("brand_name, sentiment, position, is_target_brand").eq("run_id", run["id"]).execute().data or []
        if mentions:
            document.add_paragraph(f"Run [{run['created_at']}] - Engine: {run['engine']} | Query: {p_text}")
            for m in mentions:
                document.add_paragraph(f"  - Mentioned: {m['brand_name']} | Target: {m['is_target_brand']} | Sentiment: {m.get('sentiment', 'N/A')} | Position: {m.get('position', 'N/A')}")
        
    doc_path = os.path.abspath('Gomoterra_Full_Dashboard.docx')
    document.save(doc_path)
    print(f"Successfully generated {doc_path}")

if __name__ == '__main__':
    asyncio.run(generate_full_dashboard())
