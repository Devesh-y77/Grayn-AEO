import os
from docx import Document
from app.database import get_supabase

def export_workspace_data(workspace_id):
    db = get_supabase()
    document = Document()
    document.add_heading(f'AEO Data Export - Workspace {workspace_id}', 0)
    
    # 1. Fetch Prompts
    prompts_res = db.table('aeo_prompts').select('*').eq('workspace_id', workspace_id).execute()
    prompts = {p['id']: p for p in (prompts_res.data or [])}
    
    document.add_heading('Queries / Prompts', level=1)
    if not prompts:
        document.add_paragraph('No queries found.')
    else:
        for p_id, p in prompts.items():
            document.add_paragraph(f"Prompt: {p['prompt_text']}", style='List Bullet')
            document.add_paragraph(f"Intent: {p.get('intent', 'N/A')}")
    
    # 2. Fetch Runs
    runs_res = db.table('aeo_runs').select('*').eq('workspace_id', workspace_id).order('created_at', desc=True).execute()
    runs = runs_res.data or []
    
    document.add_heading('Runs & Mentions & Citations', level=1)
    if not runs:
        document.add_paragraph('No runs found.')
        
    for run in runs:
        prompt = prompts.get(run['prompt_id'], {})
        prompt_text = prompt.get('prompt_text', 'Unknown Prompt')
        document.add_heading(f"Run: {prompt_text}", level=2)
        document.add_paragraph(f"Engine: {run['engine']} | Status: {run['status']} | Date: {run['created_at']}")
        
        if run.get('raw_response'):
            document.add_heading('Raw Response:', level=3)
            document.add_paragraph(run['raw_response'])
            
        # Mentions
        mentions_res = db.table('aeo_mentions').select('*').eq('run_id', run['id']).order('position').execute()
        mentions = mentions_res.data or []
        document.add_heading('Extracted Mentions (Competitors/Brands):', level=3)
        if mentions:
            for m in mentions:
                tgt = "(Target Brand)" if m.get('is_target_brand') else ""
                sent = f" [Sentiment: {m.get('sentiment')}]" if m.get('sentiment') else ""
                document.add_paragraph(f"{m['position']}. {m['brand_name']} {tgt}{sent}")
                if m.get('attributes'):
                    document.add_paragraph(f"Attributes: {m['attributes']}")
        else:
            document.add_paragraph('No mentions extracted.')
            
        # Citations
        citations_res = db.table('aeo_citations').select('*').eq('run_id', run['id']).execute()
        citations = citations_res.data or []
        document.add_heading('Extracted Citations:', level=3)
        if citations:
            for c in citations:
                document.add_paragraph(f"- {c.get('domain')} ({c.get('source_type')}): {c.get('url')}")
        else:
            document.add_paragraph('No citations extracted.')

    doc_path = os.path.abspath('AEO_Data_Export.docx')
    document.save(doc_path)
    print(f"Successfully exported data to {doc_path}")


if __name__ == '__main__':
    workspace_id = 'a7748191-2fc8-4596-90d6-a980308513dd'
    export_workspace_data(workspace_id)
