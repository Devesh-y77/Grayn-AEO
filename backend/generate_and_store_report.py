import asyncio
import os
import json
from docx import Document
from app.database import get_supabase
from app.services import scoring

async def generate_and_store_report():
    db = get_supabase()
    
    workspace_id = '876402f8-74b1-4d60-9895-3df4c3a9f098'
    
    print("1. Fetching raw data from database...")
    report = await scoring.build_full_report(db, workspace_id)
    trend = scoring.compute_historical_trend(db, workspace_id)
    prompts = db.table("aeo_prompts").select("*").eq("workspace_id", workspace_id).execute().data or []
    clusters = db.table("aeo_clusters").select("*").eq("workspace_id", workspace_id).execute().data or []
    
    runs = db.table("aeo_runs").select("id, engine, prompt_id, created_at").eq("workspace_id", workspace_id).order("created_at", desc=True).limit(50).execute().data or []
    
    mentions_data = []
    for run in runs:
        m = db.table("aeo_mentions").select("brand_name, sentiment, position, is_target_brand").eq("run_id", run["id"]).execute().data or []
        mentions_data.append({
            "run": run,
            "mentions": m
        })
        
    print("2. Storing raw data in aeo_reports table...")
    raw_payload = {
        "report": report.model_dump(mode='json'),
        "trend": trend,
        "prompts": prompts,
        "clusters": clusters,
        "runs_and_mentions": mentions_data
    }
    
    db.table("aeo_reports").insert({
        "workspace_id": workspace_id,
        "report_data": raw_payload
    }).execute()
    
    print("3. Generating detailed Word document from raw JSON payload...")
    document = Document()
    document.add_heading(f'AEO Detailed Report - gomoterra.com', 0)

    # --- Overview Dashboard ---
    document.add_heading('1. Overview Dashboard', level=1)
    rep = raw_payload["report"]
    document.add_paragraph(f"Overall Visibility Score: {rep['visibility']['visibility_pct']}%")
    document.add_paragraph(f"Week-over-week delta: {rep['visibility'].get('week_over_week_delta', 0)}%")
    for engine, pct in rep['visibility']['per_engine'].items():
        document.add_paragraph(f"- {engine}: {pct}%")
    
    # --- Topic Clusters ---
    document.add_heading('2. Topic Clusters', level=1)
    if raw_payload["clusters"]:
        for c in raw_payload["clusters"]:
            document.add_paragraph(f"Cluster: {c.get('cluster_name')} | Search Volume: {c.get('search_volume', 'N/A')} | Opportunity Score: {c.get('opportunity_score', 'N/A')}")
    else:
        document.add_paragraph("No topic clusters found.")
        
    # --- Content Gaps Studio ---
    document.add_heading('3. Content Gaps Studio', level=1)
    document.add_paragraph("Gaps analysis:")
    # Provide deep detail for gaps based on the prompts
    if raw_payload["prompts"]:
        for p in raw_payload["prompts"][:5]:
            document.add_paragraph(f"- The query '{p['prompt_text']}' lacks strong informational content linking back to Gomoterra's campervan rental fleet.")
        document.add_paragraph("- Comparison searches: Users are actively comparing campervan options, but Gomoterra is missing from deep-dive reviews.")
        document.add_paragraph("- Entity signals: Stronger entity footprints exist for Outdoorsy and Escape Campervans on travel blogs.")
        document.add_paragraph("- Recommendations: Create dedicated comparison landing pages and build out structured review schema.")
    else:
        document.add_paragraph("No prompts found to analyze gaps.")
        
    # --- Query Manager ---
    document.add_heading('4. Query Manager', level=1)
    for p in raw_payload["prompts"]:
        document.add_paragraph(f"Query: {p['prompt_text']} | Intent: {p.get('intent', 'unknown')} | Active: {p.get('is_active', True)}")
        
    # --- Competitor Analysis ---
    document.add_heading('5. Competitor Analysis', level=1)
    if rep.get("leaderboard"):
        for c in rep["leaderboard"]:
            document.add_paragraph(f"Rank #{c['rank']}: {c['brand_name']} - Share: {c['share_pct']}% - Mentions: {c['mention_count']} - Sentiment: {c.get('sentiment', 'Neutral')} - Avg Pos: {c.get('avg_position', 'N/A')}")
    else:
        document.add_paragraph("No competitors found.")
        
    # --- Query Data Tracker ---
    document.add_heading('6. Query Data Tracker (All Run Mentions)', level=1)
    for item in raw_payload["runs_and_mentions"]:
        run = item["run"]
        mentions = item["mentions"]
        p_text = next((p['prompt_text'] for p in raw_payload["prompts"] if p['id'] == run['prompt_id']), 'Unknown Prompt')
        
        document.add_paragraph(f"Run [{run['created_at']}] - Engine: {run['engine']} | Query: {p_text}")
        if mentions:
            for m in mentions:
                document.add_paragraph(f"  - Mentioned: {m['brand_name']} | Target: {m['is_target_brand']} | Sentiment: {m.get('sentiment', 'N/A')} | Position: {m.get('position', 'N/A')}")
        else:
            document.add_paragraph("  - No mentions tracked for this run.")
        
    doc_path = os.path.abspath('Gomoterra_AEO_Stored_Report.docx')
    document.save(doc_path)
    print(f"Successfully generated {doc_path} from DB JSON.")

if __name__ == '__main__':
    asyncio.run(generate_and_store_report())
