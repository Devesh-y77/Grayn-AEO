import asyncio
import os
from docx import Document
from app.services.discovery import run_discovery
from app.services.providers.base import get_provider, EngineType
from app.services.judge import extract_mentions_and_citations

async def generate_gomoterra_report():
    domain = "gomoterra.com"
    brand_name = "Gomoterra"
    
    document = Document()
    document.add_heading(f'AEO Detailed Report - {brand_name}', 0)
    
    document.add_paragraph("Running deep discovery and analysis...")
    
    # 1. Discovery
    print("Running discovery...")
    discovery = await run_discovery(domain, num_queries=5)
    suggested_queries = [q["text"] for q in discovery.get("suggested_queries", [])][:5]
    if not suggested_queries:
        suggested_queries = [f"Best alternatives to {domain}", f"Top services like {domain}", f"{domain} reviews"]
    
    brand_name = discovery.get("brand_name", brand_name)
    
    document.add_heading('Queries Discovered', level=1)
    for q in suggested_queries:
        document.add_paragraph(q, style='List Bullet')
        
    models = ["openai", "deepseek"]
    
    document.add_heading('Detailed Engine Runs', level=1)
    
    for query in suggested_queries:
        for engine_str in models:
            print(f"Running {engine_str} for query: {query}")
            document.add_heading(f"Query: {query} (Engine: {engine_str})", level=2)
            
            try:
                eng = EngineType(engine_str)
                provider = get_provider(eng)
                if hasattr(provider, "__aenter__"):
                    async with provider as p:
                        res = await p.query(query, location="USA")
                else:
                    res = await provider.query(query, location="USA")
                
                # Raw response
                document.add_heading('Raw Engine Response:', level=3)
                document.add_paragraph(res.raw_text)
                
                # Judge Extraction
                ext = await extract_mentions_and_citations(res.raw_text, brand_name)
                
                document.add_heading('Extracted Competitors/Brands:', level=3)
                if ext.mentions:
                    for m in ext.mentions:
                        tgt = "(Target Brand)" if m.is_target_brand else ""
                        sent = f" [Sentiment: {m.sentiment}]" if m.sentiment else ""
                        document.add_paragraph(f"{m.position}. {m.brand_name} {tgt}{sent}")
                else:
                    document.add_paragraph("No brands extracted.")
                    
                document.add_heading('Extracted Citations:', level=3)
                if ext.citations:
                    for c in ext.citations:
                        document.add_paragraph(f"- {c.domain} ({c.source_type}): {c.url}")
                else:
                    document.add_paragraph("No citations extracted.")
                    
            except Exception as e:
                document.add_paragraph(f"Error during execution: {str(e)}")

    doc_path = os.path.abspath('Gomoterra_AEO_Report.docx')
    document.save(doc_path)
    print(f"Successfully generated full report at {doc_path}")

if __name__ == '__main__':
    asyncio.run(generate_gomoterra_report())
